"""
File Name: generate_docs.py
Author: Senthilnathan Karuppaiah
Date: 03-May-2026
Description: One-time script to generate synthetic documents from articles_100.train.
             Each row in the training file is converted into exactly one output document.
             File types are distributed evenly across the 100 rows by cycling through
             .docx → .pdf → .md in sequence. File names are derived from distinctive
             content words found in the article text, ensuring each name is unique and
             semantically meaningful.

Note: Run this script once from the minhash-doc-similarity directory using the local
      virtual environment:
          .venv/bin/python generate_docs.py
      Output is written to data/input/. Re-running will overwrite existing files.

Requirements:
- python-docx
- reportlab
"""

import logging
import os
import re
import textwrap
from datetime import datetime
from pathlib import Path

import structlog
from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


# ── logging ────────────────────────────────────────────────────────────────────

script_name = os.path.splitext(os.path.basename(__file__))[0]
log_filename = f"{script_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

logging.basicConfig(filename=log_filename, level=logging.INFO, format="%(message)s")
structlog.configure(
    logger_factory=structlog.stdlib.LoggerFactory(),
    processors=[
        structlog.stdlib.filter_by_level,
        structlog.stdlib.add_logger_name,
        structlog.stdlib.add_log_level,
        structlog.stdlib.PositionalArgumentsFormatter(),
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.processors.JSONRenderer(),
    ],
)
log = structlog.get_logger()


# ── constants ──────────────────────────────────────────────────────────────────

TRAIN_FILE = Path(__file__).parent.parent / "MinHash" / "data" / "articles_100.train"
OUTPUT_DIR = Path(__file__).parent / "data" / "input"

# Cycles over the 100 rows assigning one format each
FILE_TYPE_CYCLE = ["docx", "pdf", "md"]

# Common English stopwords to skip when building content slugs
STOPWORDS = {
    "a", "an", "the", "and", "or", "but", "in", "on", "at", "to", "for",
    "of", "with", "by", "from", "as", "is", "was", "are", "were", "be",
    "been", "being", "have", "has", "had", "do", "does", "did", "will",
    "would", "could", "should", "may", "might", "shall", "not", "no",
    "it", "its", "he", "she", "they", "we", "i", "his", "her", "their",
    "that", "this", "which", "who", "when", "where", "how", "what", "after",
    "before", "into", "over", "also", "said", "says", "say", "s", "about",
    "up", "out", "if", "than", "then", "so", "more", "two", "one", "three",
    "new", "first", "last", "other", "such", "there", "here", "just",
}


# ── helpers ────────────────────────────────────────────────────────────────────

def load_training_rows(file_path: Path) -> list[tuple[str, str]]:
    """
    Read the training file and return a list of (doc_id, text) tuples.
    Each line is expected to start with a whitespace-separated token ID
    followed by the article body.

    Args:
        file_path: Absolute path to the .train file.

    Returns:
        List of (doc_id, text) tuples, one per non-empty line.
    """
    rows = []
    with open(file_path, encoding="utf-8") as fh:
        for line in fh:
            line = line.strip()
            if not line:
                continue
            parts = line.split(" ", 1)
            doc_id = parts[0] if len(parts) == 2 else f"doc{len(rows) + 1}"
            text = parts[1] if len(parts) == 2 else parts[0]
            rows.append((doc_id, text))
    log.info("training_rows_loaded", count=len(rows), source=str(file_path))
    return rows


def extract_title(text: str) -> str:
    """
    Derive a human-readable document title from the first sentence of the text.
    Truncates to 80 characters to keep titles concise.

    Args:
        text: Full article body string.

    Returns:
        A short title string.
    """
    first_sentence = re.split(r"(?<=[.!?])\s+", text.strip())[0]
    return first_sentence[:80].rstrip(".!?,") if first_sentence else "News Article"


def build_content_slug(doc_id: str, text: str, seen_slugs: set[str]) -> str:
    """
    Build a unique, content-derived filename slug from the article text.
    Extracts the first three meaningful (non-stopword, alphabetic) words
    from the text and combines them with the doc_id for uniqueness.

    Args:
        doc_id:      Source document identifier from the training file (e.g. 't980').
        text:        Full article body string.
        seen_slugs:  Set of slugs already used; updated in-place on each call.

    Returns:
        A lowercase underscore-separated slug string, e.g. 'zambian_police_ritual_t980'.
    """
    words = re.findall(r"[a-zA-Z]{3,}", text)
    keywords = [w.lower() for w in words if w.lower() not in STOPWORDS][:3]
    base = "_".join(keywords) + f"_{doc_id}" if keywords else doc_id

    # Guarantee uniqueness by appending a counter if the slug has been used
    slug = base
    counter = 2
    while slug in seen_slugs:
        slug = f"{base}_{counter}"
        counter += 1

    seen_slugs.add(slug)
    return slug


# ── document writers ───────────────────────────────────────────────────────────

def write_docx(output_path: Path, title: str, body: str) -> None:
    """
    Write article content as a formatted Microsoft Word (.docx) document.

    Args:
        output_path: Full destination file path including filename and extension.
        title:       Document heading text.
        body:        Article body to be written as wrapped paragraphs.

    Returns:
        None
    """
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in textwrap.wrap(body, width=100):
        doc.add_paragraph(paragraph)
    doc.save(str(output_path))
    log.info("docx_written", path=str(output_path))


def write_pdf(output_path: Path, title: str, body: str) -> None:
    """
    Write article content as a PDF document using ReportLab.

    Args:
        output_path: Full destination file path including filename and extension.
        title:       Document title rendered using the 'Title' paragraph style.
        body:        Article body rendered as wrapped 'BodyText' paragraphs.

    Returns:
        None
    """
    styles = getSampleStyleSheet()
    story = [
        Paragraph(title, styles["Title"]),
        Spacer(1, 12),
    ]
    for paragraph in textwrap.wrap(body, width=120):
        story.append(Paragraph(paragraph, styles["BodyText"]))
        story.append(Spacer(1, 6))
    pdf_doc = SimpleDocTemplate(str(output_path), pagesize=LETTER)
    pdf_doc.build(story)
    log.info("pdf_written", path=str(output_path))


def write_md(output_path: Path, title: str, body: str) -> None:
    """
    Write article content as a Markdown (.md) file.

    Args:
        output_path: Full destination file path including filename and extension.
        title:       Rendered as a level-1 Markdown heading.
        body:        Article body written as plain text below the heading.

    Returns:
        None
    """
    with open(output_path, "w", encoding="utf-8") as fh:
        fh.write(f"# {title}\n\n")
        fh.write(body.strip())
        fh.write("\n")
    log.info("md_written", path=str(output_path))


# ── dispatcher ─────────────────────────────────────────────────────────────────

WRITERS = {
    "docx": write_docx,
    "pdf":  write_pdf,
    "md":   write_md,
}


def generate_document(idx: int, doc_id: str, text: str, seen_slugs: set[str]) -> str:
    """
    Generate a single output document for one training-file row.
    The file type is determined by cycling through docx → pdf → md
    based on the row index.

    Args:
        idx:         1-based row index used to select the file type from the cycle.
        doc_id:      Source document identifier from the training file.
        text:        Article body text.
        seen_slugs:  Shared set of already-used slugs to enforce uniqueness.

    Returns:
        The filename (without directory path) of the generated document.
    """
    file_type = FILE_TYPE_CYCLE[(idx - 1) % len(FILE_TYPE_CYCLE)]
    slug = build_content_slug(doc_id, text, seen_slugs)
    title = extract_title(text)
    output_path = OUTPUT_DIR / f"{slug}.{file_type}"

    WRITERS[file_type](output_path, title, text)
    return output_path.name


# ── main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    """
    Entry point. Reads the training file, generates one document per row,
    and logs a summary on completion.
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = load_training_rows(TRAIN_FILE)

    log.info("generation_started", total_rows=len(rows), output_dir=str(OUTPUT_DIR))
    print(f"Generating {len(rows)} documents → {OUTPUT_DIR}\n")

    seen_slugs: set[str] = set()

    for idx, (doc_id, text) in enumerate(rows, start=1):
        filename = generate_document(idx, doc_id, text, seen_slugs)
        print(f"  [{idx:03d}]  {filename}")

    log.info("generation_complete", total_files=len(rows), output_dir=str(OUTPUT_DIR))
    print(f"\nDone. {len(rows)} files written to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
